#!/usr/bin/env python3
"""Extract visible DOCX paragraphs and core properties into JSON."""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()
    try:
        from docx import Document
    except ImportError as exc:
        raise SystemExit("python-docx is required") from exc
    document = Document(args.docx)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    tables = [[[cell.text for cell in row.cells] for row in table.rows] for table in document.tables]
    props = document.core_properties
    result = {
        "paragraphs": paragraphs,
        "tables": tables,
        "metadata": {"title": props.title or "", "author": props.author or "", "subject": props.subject or ""},
    }
    rendered = json.dumps(result, ensure_ascii=False, indent=2)
    if args.output:
        args.output.write_text(rendered + "\n", encoding="utf-8")
    else:
        print(rendered)
    return 0


if __name__ == "__main__":
    sys.exit(main())
