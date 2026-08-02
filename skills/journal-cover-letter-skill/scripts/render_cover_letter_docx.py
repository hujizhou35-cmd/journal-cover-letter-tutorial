#!/usr/bin/env python3
"""Render a validated structured payload as a privacy-scrubbed DOCX."""

from __future__ import annotations

import argparse
import json
import re
import sys
import tempfile
import zipfile
from pathlib import Path


def build(payload: dict, output: Path) -> None:
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise SystemExit("python-docx is required") from exc

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.10

    if payload.get("date"):
        document.add_paragraph(payload["date"])
    if payload.get("recipient_block"):
        for line in payload["recipient_block"]:
            document.add_paragraph(line)
    document.add_paragraph(payload["salutation"])
    for text in payload["paragraphs"]:
        paragraph = document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph(payload["signoff"])
    document.add_paragraph(payload["corresponding_author"])
    if payload.get("affiliation"):
        document.add_paragraph(payload["affiliation"])
    if payload.get("contact"):
        document.add_paragraph(payload["contact"])

    props = document.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.title = "Academic Journal Cover Letter"
    props.subject = ""
    props.comments = ""
    props.keywords = ""
    props.category = ""
    document.save(output)
    scrub_metadata(output)


def scrub_metadata(path: Path) -> None:
    """Remove core authorship, custom properties, and Word rsid attributes."""
    namespaces = {
        "dc_creator": re.compile(rb"(<dc:creator[^>]*>).*?(</dc:creator>)", re.S),
        "last_modified": re.compile(rb"(<cp:lastModifiedBy[^>]*>).*?(</cp:lastModifiedBy>)", re.S),
        "rsid": re.compile(rb"\s+w:rsid[A-Za-z]*=\"[^\"]*\""),
    }
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=path.parent) as handle:
        temp_path = Path(handle.name)
    try:
        with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
            for info in source.infolist():
                name = info.filename
                if name == "docProps/custom.xml":
                    continue
                content = source.read(name)
                if name == "docProps/core.xml":
                    content = namespaces["dc_creator"].sub(rb"\1\2", content)
                    content = namespaces["last_modified"].sub(rb"\1\2", content)
                elif name.startswith("word/") and name.endswith(".xml"):
                    content = namespaces["rsid"].sub(b"", content)
                elif name == "_rels/.rels":
                    content = re.sub(rb"<Relationship\b[^>]*Target=\"docProps/custom\.xml\"[^>]*/>", b"", content)
                elif name == "[Content_Types].xml":
                    content = re.sub(rb"<Override\b[^>]*PartName=\"/docProps/custom\.xml\"[^>]*/>", b"", content)
                target.writestr(info, content)
        temp_path.replace(path)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("payload", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    payload = json.loads(args.payload.read_text(encoding="utf-8"))
    from validate_payload import validate

    errors = validate(payload)
    if errors:
        print(json.dumps({"valid": False, "errors": errors}, indent=2), file=sys.stderr)
        return 1
    build(payload, args.output)
    print(args.output)
    return 0


if __name__ == "__main__":
    sys.exit(main())
